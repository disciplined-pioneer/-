import os
import math
import locale
import subprocess
from datetime import datetime, timedelta
from utils.check import get_last_check_id

from docx import Document
from docx.shared import Pt
from num2words import num2words
from docx.oxml import OxmlElement, ns


MAX_PARTICIPANTS = 10  # Максимальное количество участников


# Функция для преобразования суммы в текст
def convert_number_to_text(number):

    # Выводим сумму как текст
    locale.setlocale(locale.LC_ALL, 'ru_RU.UTF-8')
    currency_text = locale.currency(number, grouping=True).replace('₽', '')
    return currency_text.replace('\xa0', ' ')


# Функция для перевода суммы в нужный формат
def process_data(data):

    # Получение значений
    sum_value = float(data.get("answers_check")['sum']) 
    date_input = data.get("answers_check")['date']

    # Q9: Целая часть от sum
    Q9 = int(sum_value) 
    
    # W9: Копейки от sum
    W9 = round((sum_value - Q9) * 100)

    # Определяем формат входящей даты
    if isinstance(date_input, datetime): 
        date_obj = date_input
    else:
        try:
            date_obj = datetime.strptime(date_input, '%Y%m%dT%H%M') 
        except ValueError:
            date_obj = datetime.strptime(date_input, '%d.%m.%Y %H:%M') 

    # I13: Дата в формате 'ДД.ММ.ГГ'
    I13 = date_obj.strftime('%d.%m.%y') 

    # I33: Сумма (просто отображаем sum_value)
    I33 = sum_value

    # P23: Форматированное слово "Расход" + месяц и год
    months = {
        'January': 'Январь', 'February': 'Февраль', 'March': 'Март', 'April': 'Апрель', 
        'May': 'Май', 'June': 'Июнь', 'July': 'Июль', 'August': 'Август', 'September': 'Сентябрь', 
        'October': 'Октябрь', 'November': 'Ноябрь', 'December': 'Декабрь'
    }

    month_eng = date_obj.strftime('%B')  # Получаем название месяца на английском
    P23 = f"Расход {months.get(month_eng, month_eng)} {date_obj.strftime('%Y')}"

    # I39: Преобразуем сумму в текст
    rub, kop = divmod(sum_value, 1) 
    kop = round(kop * 100)
    I39 = f"{convert_number_to_text(int(rub))} рублей {kop} копеек ({int(rub)} руб. {kop} коп.)"

    # Результат в виде словаря
    return {
        'Q9': Q9,
        'W9': W9,
        'I13': I13,
        'I33': I33,
        'I35': I33,
        'I39': I39,
        'P23': P23
    }


questions_event = {
    "company_meeting": ('Введите данные для заполнения шаблона:\n\n💼 Компания, с которой планируется втреча:'),
    "event_location": ('🏢 Место проведения:'),
    "meeting_theme": ("📝 Тема собрания. Введите тему: "),
    "guest_name": ("Пожалуйста, введите ФИО приглашенного участника: "),
    "guest_workplace": ("Пожалуйста, выберите место работы для участника: ")
}


# Функция для перевода числа в текст на русском языке
def num_to_text(n):
    return num2words(n, lang='ru').replace(",", "").replace("-", " ")

# Словарь для месяцев на русском языке
months = [
    "января", "февраля", "марта", "апреля", "мая", "июня", 
    "июля", "августа", "сентября", "октября", "ноября", "декабря"
]

# Функция для обработки документа и замены данных
async def process_document(doc_path, data, user, file_path):

    # Загружаем существующий документ
    doc = Document(doc_path)

    # Функция для вычисления времени с отнятием 2 часов
    def subtract_hours_from_time(date_obj, hours=2):
        if isinstance(date_obj, datetime):
            new_time = date_obj - timedelta(hours=hours)
            return new_time.strftime("%H:%M")
        return ""

    # Получаем даты
    date_compilation = data.get('answers_check', {}).get('date')
    if isinstance(date_compilation, datetime):
        date_compilation_str = f"{(date_compilation - timedelta(days=7)).day} {months[(date_compilation - timedelta(days=7)).month - 1]} {date_compilation.year} г."
        period_implem = f"{date_compilation.day} {months[date_compilation.month - 1]} {date_compilation.year} г."
    else:
        date_compilation_str = ""
        period_implem = ""

    # Заменяем сумму (округляем до тысячи)
    sum_check = data.get('answers_check', {}).get('sum')
    if isinstance(sum_check, (int, float)):
        sum_check = math.ceil(sum_check / 1000) * 1000
        sum_check_str = f"{sum_check:,.2f}".replace(",", " ").replace('.', ',') + " рублей"
        sum_check_text = num_to_text(sum_check).capitalize() + " рублей 00 копеек"
    else:
        sum_check_str = ""
        sum_check_text = ""

    # Получаем время (если дата указана)
    start_time = subtract_hours_from_time(date_compilation)
    middle_start_time = subtract_hours_from_time(date_compilation, hours=2)
    middle_end_time = subtract_hours_from_time(date_compilation, hours=1)
    end_time = subtract_hours_from_time(date_compilation, hours=2)
    last_check_id = await get_last_check_id()

    # Данные для замены
    placeholders = {
        "{date_compilation}": date_compilation_str,
        "{period_implem}": period_implem,
        "{event_location}": data.get('answers', {}).get('event_location', ""),
        "{job_title}": user.get('subdivision', ""),
        "{snp}": user.get('full_name', ""),
        "{meeting_theme}": data.get('answers', {}).get('meeting_theme', ""),
        "{sum_check}": sum_check_str,
        "{sum_check_text}": sum_check_text,
        "{start_time}": start_time,
        "{middle_start_time}": middle_start_time,
        "{middle_end_time}": middle_end_time,
        "{end_time}": end_time,
        "{name_holiday}": data.get('answers', {}).get('event', ""),
        "{gifts_text}": "\n".join(f"{gift}" for gift in data.get('answers', {}).get('gifts', "")),
        "{selected_drug}": data.get('selected_drug', ''),
        "{company_meeting}": data.get('answers', {}).get('company_meeting', ""),
        "{name_gift}": data.get('answers', {}).get('name_gift', ""),
        "{number_check}": str(last_check_id + 1)
    }

    # Проходим по всем параграфам и заменяем текст
    for para in doc.paragraphs:
        for key, value in placeholders.items():
            if key in para.text:
                para.text = para.text.replace(key, value)

    # Проходим по всем таблицам и заменяем текст в ячейках
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in placeholders.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)

    # Сохраняем изменения
    doc.save(file_path)

    print("Данные в файле Word заменены!")


# Функция для установки границ ячейки
def set_cell_border(cell):
    tc = cell._element
    tc_pr = tc.get_or_add_tcPr()
    for border_name in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f'w:{border_name}')
        border.set(ns.qn('w:val'), 'single')
        border.set(ns.qn('w:sz'), '6')
        border.set(ns.qn('w:space'), '0')
        border.set(ns.qn('w:color'), '000000')
        tc_pr.append(border)


# Функция для добавления новой строки в таблицу
def add_row_with_borders(table, data):
    new_row = table.add_row()  # Добавляем новую строку
    num_cols = len(table.rows[0].cells)  # Количество колонок в таблице

    # Если данных меньше, чем колонок — заполняем только доступные
    for col_idx, text in enumerate(data):
        if col_idx >= num_cols:
            print(f"⚠ Пропущен элемент '{text}', так как в таблице {num_cols} колонок")
            continue

        cell = new_row.cells[col_idx]
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(str(text))  # Преобразуем в строку (на случай чисел)
        run.font.size = Pt(11)  # Устанавливаем размер шрифта
        set_cell_border(cell)  # Устанавливаем границы


# Функция для изменения данных последней строки таблицы
def update_last_row(table, data):
    last_row = table.rows[-1]  # Получаем последнюю строку
    for col_idx, text in enumerate(data):
        cell = last_row.cells[col_idx]
        cell.text = text  # Заменяем текст в ячейке
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(11)  # Устанавливаем размер шрифта
        set_cell_border(cell)  # Устанавливаем границы


# Конвертация word в pdf
def convert_docx_to_pdf(input_file, output_file):

    # Проверка, что входной файл существует
    if not os.path.isfile(input_file):
        print(f"Файл {input_file} не найден.")
        return

    # Папка, где будет сохранён файл
    output_dir = os.path.dirname(output_file)
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # Строим команду для LibreOffice
    command = [
        "libreoffice",
        "--headless",
        "--convert-to", 
        "pdf",
        input_file,
        "--outdir", 
        output_dir
    ]
    
    try:
        # Запуск команды
        subprocess.run(command, check=True)

        # Получаем имя выходного файла
        output_pdf = os.path.join(output_dir, os.path.basename(input_file).replace(".docx", ".pdf"))

        # Переименовываем файл
        if output_pdf != output_file:
            os.rename(output_pdf, output_file)

        print(f"Файл успешно конвертирован в PDF: {output_file}")
    except subprocess.CalledProcessError as e:
        print(f"Ошибка при конвертации: {e}")

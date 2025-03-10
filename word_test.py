from docx import Document
from datetime import datetime, timedelta
import math
from num2words import num2words

# Функция для перевода числа в текст на русском языке
def num_to_text(n):
    return num2words(n, lang='ru').replace(",", "").replace("-", " ")

# Словарь для месяцев на русском языке
months = [
    "января", "февраля", "марта", "апреля", "мая", "июня", 
    "июля", "августа", "сентября", "октября", "ноября", "декабря"
]

# Функция для обработки документа и замены данных
def process_document(doc_path, data, user):
    # Загружаем существующий документ
    doc = Document(doc_path)

    # Функция для вычисления времени с отнятием 2 часов
    def subtract_hours_from_time(date_obj, hours=2):
        if isinstance(date_obj, datetime):
            new_time = date_obj - timedelta(hours=hours)
            return new_time.strftime("%H:%M")
        return ""

    # Получаем даты
    date_compilation = data.get('answers_check', {}).get('date')
    if isinstance(date_compilation, datetime):
        date_compilation_str = f"{(date_compilation - timedelta(days=7)).day} {months[(date_compilation - timedelta(days=7)).month - 1]} {date_compilation.year} г."
        period_implem = f"{date_compilation.day} {months[date_compilation.month - 1]} {date_compilation.year} г."
    else:
        date_compilation_str = ""
        period_implem = ""

    # Заменяем сумму (округляем до тысячи)
    sum_check = data.get('answers_check', {}).get('sum')
    if isinstance(sum_check, (int, float)):
        sum_check = math.ceil(sum_check / 1000) * 1000
        sum_check_str = f"{sum_check:,.2f}".replace(",", " ").replace('.', ',') + " рублей"
        sum_check_text = num_to_text(sum_check).capitalize() + " рублей 00 копеек"
    else:
        sum_check_str = ""
        sum_check_text = ""

    # Получаем время (если дата указана)
    start_time = subtract_hours_from_time(date_compilation)
    middle_start_time = subtract_hours_from_time(date_compilation, hours=2)
    middle_end_time = subtract_hours_from_time(date_compilation, hours=1)
    end_time = subtract_hours_from_time(date_compilation, hours=2)

    # Данные для замены
    placeholders = {
        "{date_compilation}": date_compilation_str,
        "{period_implem}": period_implem,
        "{location}": data.get('answers_check', {}).get('event_location', ""),
        "{job_title}": user.get('subdivision', ""),
        "{snp}": user.get('full_name', ""),
        "{topic}": data.get('answers_check', {}).get('meeting_theme', ""),
        "{sum_check}": sum_check_str,
        "{sum_check_text}": sum_check_text,
        "{start_time}": start_time,
        "{middle_start_time}": middle_start_time,
        "{middle_end_time}": middle_end_time,
        "{end_time}": end_time,
        "{name_holiday}": data.get('answers_check', {}).get('event', "")
    }

    # Проходим по всем параграфам и заменяем текст
    for para in doc.paragraphs:
        for key, value in placeholders.items():
            if key in para.text:
                para.text = para.text.replace(key, value)

    # Проходим по всем таблицам и заменяем текст в ячейках
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in placeholders.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)

    # Сохраняем изменения
    doc.save("data/output.docx")

    print("Данные в файле Word заменены!")

# Пример использования функции
data = {
    'answers_check': {
        'date': datetime(2024, 12, 29, 14, 50),  # Пример с временем
        'event_location': 'Москва',
        'meeting_theme': 'Обсуждение важного вопроса',
        'sum': 17500.45,
        'event': 'Новый год'
    }
}

user = {
    'subdivision': 'Отдел разработки',
    'full_name': 'Иванов Иван Иванович'
}

doc_path = "data/present_test.docx"

process_document(doc_path, data, user)
